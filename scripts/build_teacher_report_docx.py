from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT = Path(__file__).resolve().parents[1]
OUT_DIR = PROJECT / "reports"
ASSET_DIR = OUT_DIR / "_teacher_report_assets"
OUTPUT = OUT_DIR / "视觉跌倒检测模型运行测试与系统对接阶段报告.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
GREEN = "177245"
GOLD = "9A6700"
RED = "9B1C1C"
WHITE = "FFFFFF"


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    total = sum(widths)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, fld_sep, text, fld_end])
    set_run_font(run, size=9, color=MUTED)


def add_paragraph(doc, text="", *, style=None, bold_prefix=None, color=None, align=None,
                  before=0, after=6, line=1.1, keep=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_with_next = keep
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True, color=color)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, color=color)
    return p


def add_bullet(doc, text, level=0):
    p = add_paragraph(doc, text, style="List Bullet", after=5, line=1.167)
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    return p


def add_number(doc, text):
    p = add_paragraph(doc, text, style="List Number", after=5, line=1.167)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, bold=True)
    return p


def add_callout(doc, label, text, fill=PALE_BLUE, accent=BLUE):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "20")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), accent)
    borders.append(left)
    p_pr.append(borders)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.1
    r1 = p.add_run(label + " ")
    set_run_font(r1, bold=True, color=accent)
    r2 = p.add_run(text)
    set_run_font(r2, color=INK)
    return p


def add_picture_with_alt(doc, path, width, alt_text):
    shape = doc.add_picture(str(path), width=width)
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt_text)
    doc_pr.set("title", alt_text)
    return shape


def add_table(doc, headers, rows, widths, numeric_cols=()):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(str(header))
        set_run_font(r, size=9.5, bold=True, color=INK)
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in numeric_cols else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(value))
            set_run_font(r, size=9.5)
    add_paragraph(doc, "", after=3)
    return table


def make_bar_chart(path):
    labels = [
        "YOLO-Pose+ByteTrack\n+ST-GCN++",
        "RTMPose\n+ST-GCN++",
        "RTMPose+ByteTrack\n+ST-GCN++",
        "RTMPose\n+PoseC3D-style",
        "RTMPose+ByteTrack v2\n+ST-GCN++",
        "RTMPose\n+CTR-GCN",
        "YOLO-Pose+ByteTrack\n+CTR-GCN",
        "YOLO-Pose+ByteTrack\n+PoseC3D-style",
    ]
    values = [89.41, 86.83, 86.29, 86.29, 84.32, 83.15, 83.13, 82.56]
    colors = ["#2E74B5", "#177245", "#7895B2", "#9BB3C8", "#B2C3D3", "#D0D9E2", "#D9E1E8", "#E4E9EE"]
    image = Image.new("RGB", (1900, 1020), "white")
    draw = ImageDraw.Draw(image)
    font_path = r"C:\Windows\Fonts\arial.ttf"
    bold_path = r"C:\Windows\Fonts\arialbd.ttf"
    title_font = ImageFont.truetype(bold_path, 42)
    label_font = ImageFont.truetype(font_path, 27)
    value_font = ImageFont.truetype(bold_path, 27)
    axis_font = ImageFont.truetype(font_path, 25)
    draw.text((950, 45), "GMDCSA24 subject-isolated test: leading routes",
              font=title_font, fill="#0B2545", anchor="ma")
    left, right, top, row_h = 650, 1760, 145, 100
    min_v, max_v = 75.0, 91.0
    for tick in range(76, 92, 2):
        x = left + int((tick - min_v) / (max_v - min_v) * (right - left))
        draw.line((x, top, x, top + row_h * len(labels)), fill="#E3E7EB", width=2)
        draw.text((x, top + row_h * len(labels) + 18), str(tick),
                  font=axis_font, fill="#667085", anchor="ma")
    for idx, (label, value, color) in enumerate(zip(labels, values, colors)):
        y = top + idx * row_h
        line1, *rest = label.split("\n")
        draw.text((left - 22, y + 28), line1, font=label_font, fill="#344054", anchor="ra")
        if rest:
            draw.text((left - 22, y + 62), rest[0], font=label_font, fill="#344054", anchor="ra")
        width = int((value - min_v) / (max_v - min_v) * (right - left))
        draw.rounded_rectangle((left, y + 22, left + width, y + 76), radius=10, fill=color)
        draw.text((left + width + 14, y + 49), f"{value:.2f}%",
                  font=value_font, fill="#0B2545", anchor="lm")
    draw.text(((left + right) // 2, 990), "Balanced Accuracy (%)",
              font=axis_font, fill="#667085", anchor="ms")
    image.save(path, quality=95)


def make_architecture_chart(path):
    image = Image.new("RGB", (1980, 480), "white")
    draw = ImageDraw.Draw(image)
    font_path = r"C:\Windows\Fonts\arial.ttf"
    bold_path = r"C:\Windows\Fonts\arialbd.ttf"
    box_font = ImageFont.truetype(bold_path, 29)
    note_font = ImageFont.truetype(font_path, 25)
    labels = [
        "Camera / RTSP\nor video",
        "RTMPose\nCOCO-17",
        "64-frame\nsliding window",
        "ST-GCN++\nprobability",
        "State machine\nand event API",
    ]
    x_positions = [35, 430, 825, 1220, 1615]
    box_w, box_h, y = 300, 175, 105
    for x, label in zip(x_positions, labels):
        draw.rounded_rectangle((x, y, x + box_w, y + box_h), radius=18,
                               fill="#E8EEF5", outline="#2E74B5", width=4)
        lines = label.split("\n")
        draw.text((x + box_w / 2, y + 65), lines[0], font=box_font, fill="#0B2545", anchor="mm")
        draw.text((x + box_w / 2, y + 112), lines[1], font=box_font, fill="#0B2545", anchor="mm")
    for x in [350, 745, 1140, 1535]:
        draw.line((x, 192, x + 55, 192), fill="#667085", width=5)
        draw.polygon([(x + 55, 192), (x + 38, 181), (x + 38, 203)], fill="#667085")
    note = "Output: CONFIRMED fall event, probability, camera/track ID, timestamp and evidence clip"
    draw.text((990, 382), note, font=note_font, fill="#667085", anchor="mm")
    image.save(path, quality=95)


def configure_document(doc):
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    heading_tokens = {
        1: (16, BLUE, 16, 8),
        2: (13, BLUE, 12, 6),
        3: (12, DARK_BLUE, 8, 4),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.167
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("视觉跌倒检测模型运行测试与系统对接阶段报告")
    set_run_font(hr, size=9, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run("第 ")
    set_run_font(fr, size=9, color=MUTED)
    add_page_field(fp)
    fr2 = fp.add_run(" 页")
    set_run_font(fr2, size=9, color=MUTED)


def build_document():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    bar_chart = ASSET_DIR / "top_routes_ba.png"
    architecture = ASSET_DIR / "system_architecture.png"
    make_bar_chart(bar_chart)
    make_architecture_chart(architecture)

    doc = Document()
    configure_document(doc)

    # Editorial cover pattern, implemented within the standard business brief preset.
    add_paragraph(doc, "", after=70)
    p = add_paragraph(doc, "阶段技术报告", align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    set_run_font(p.runs[0], size=11, bold=True, color=GOLD)
    p = add_paragraph(doc, "视觉跌倒检测模型运行测试\n与系统对接阶段报告",
                      align=WD_ALIGN_PARAGRAPH.CENTER, after=12, line=1.05)
    for run in p.runs:
        set_run_font(run, size=28, bold=True, color=INK)
    p = add_paragraph(doc, "从模型调研、实际训练测试到部署对齐验证",
                      align=WD_ALIGN_PARAGRAPH.CENTER, after=40)
    set_run_font(p.runs[0], size=14, color=DARK_BLUE)
    add_paragraph(doc, "项目组", align=WD_ALIGN_PARAGRAPH.CENTER, after=5, color=MUTED)
    add_paragraph(doc, "2026 年 7 月 24 日", align=WD_ALIGN_PARAGRAPH.CENTER, after=6, color=MUTED)
    add_paragraph(doc, "报告用途：向指导教师汇报模型实测效果，并作为后续系统对接讨论依据",
                  align=WD_ALIGN_PARAGRAPH.CENTER, before=90, after=0, color=MUTED)
    doc.add_page_break()

    add_heading(doc, "摘要", 1)
    add_paragraph(
        doc,
        "本阶段工作已从“仅调研模型”推进到“统一协议下实际运行、训练、测试并形成可运行系统原型”。"
        "项目最初重点比较 YOLO-Pose + ST-GCN++、RTMPose + ST-GCN++ 和 "
        "YOLO-Pose + CTR-GCN，随后扩展为 7 个姿态/跟踪前端与 3 个时序分类器的 "
        "21 路正交网格，并增加 2 条 ByteTrack 消融路线，共完成 23 条路线的四折 300 轮训练。"
        "综合主体隔离内部测试与 MCFD 外部测试，RTMPose + ST-GCN++ 的跨场景表现最均衡。"
        "为解决训练阶段均匀采样与系统滑窗输入不一致的问题，又重新构建 1548 个连续窗口并训练部署对齐模型。"
        "当前系统在固定 8 段整视频冒烟测试中实现 4/4 跌倒检出和 4/4 ADL 无误报。"
    )
    add_callout(
        doc,
        "阶段结论：",
        "当前推荐将 RTMPose + ST-GCN++ 作为系统基线；YOLO-Pose + ByteTrack + ST-GCN++ "
        "保留为内部高分研究路线。下一阶段的重点不是继续堆叠模型，而是完成设备接入、多人跟踪、"
        "现场数据采集以及事件级误报率、漏报率和报警延迟评价。",
    )

    add_heading(doc, "1. 教师要求与本阶段响应", 1)
    add_paragraph(
        doc,
        "教师提出的核心要求是：必须选择模型实际运行并反馈测试效果，不能只停留在文献调研；"
        "在看到模型效果之后，再共同讨论如何接入系统。围绕这一要求，本阶段完成了以下闭环："
    )
    for text in [
        "确定候选路线并统一数据、骨架格式、训练轮数和评价指标。",
        "实际提取姿态、训练时序模型、保存各折最优权重并生成学习曲线。",
        "在主体隔离内部测试和未参与训练的 MCFD 外部数据上比较模型。",
        "将筛选出的 RTMPose + ST-GCN++ 接入预录视频流水线，运行整视频演示。",
        "针对训练/部署分布不一致重新构建滑窗数据并训练系统权重。",
        "形成可供后续摄像头、RTSP、报警平台和业务系统对接的模块化方案。",
    ]:
        add_number(doc, text)

    add_heading(doc, "2. 实验对象与统一协议", 1)
    add_heading(doc, "2.1 实测模型范围", 2)
    add_paragraph(
        doc,
        "核心实验覆盖 7 个姿态或跟踪前端与 3 个时序分类器。ByteTrack 是人物关联器，"
        "不是跌倒分类器，因此作为姿态前端变体参加组合实验。"
    )
    add_table(
        doc,
        ["模块", "实际运行对象", "作用"],
        [
            ["姿态/跟踪前端", "RTMPose、YOLO-Pose、YOLO-Pose+ByteTrack、RTMO、Hourglass52、OpenPose、AlphaPose", "从视频帧提取人体关键点或稳定主轨"],
            ["时序分类器", "ST-GCN++、CTR-GCN、PoseC3D-style", "根据 64 帧骨架序列判断 Fall/ADL"],
            ["附加消融", "RTMPose+ByteTrack、RTMPose+ByteTrack v2", "分析跟踪丢失、回接和插值的影响"],
        ],
        [1800, 4680, 2880],
    )
    add_paragraph(
        doc,
        "其中 PoseC3D-style 是项目内实现：将归一化骨架渲染为时空热图体，再由 3D 残差卷积分类；"
        "它采用 PoseC3D 的核心表示思想，但不等同于 MMAction2 官方完整配置。",
        color=MUTED,
    )

    add_heading(doc, "2.2 数据集与划分", 2)
    add_table(
        doc,
        ["数据集", "用途", "规模与划分"],
        [
            ["GMDCSA24", "训练、验证、主体隔离内部测试", "160 个视频；81 ADL、79 Fall；4 名受试者四折隔离"],
            ["MCFD", "未见场景外部测试", "552 个标注片段；cam1 校准，cam3 开发观察，cam2/4/5/6/7/8 共 415 段正式跨视角测试"],
            ["GMDCSA24 连续滑窗", "系统部署对齐训练", "1548 个窗口；1071 负类、477 正类；36 个边界窗口忽略"],
        ],
        [1800, 2640, 4920],
    )
    add_bullet(doc, "统一骨架：COCO-17；通道为归一化 x、y 和置信度。")
    add_bullet(doc, "统一输入： [N,C,T,V,M] = [样本,3,64,17,1]。")
    add_bullet(doc, "主体隔离四折：每折使用两个 Subject 训练、一个验证、一个测试；合并测试覆盖全部样本一次。")
    add_bullet(doc, "外部测试数据不参与训练，用于观察跨数据集、跨视角与动作定义差异。")

    add_heading(doc, "2.3 训练与评价设置", 2)
    add_table(
        doc,
        ["项目", "23 路模型比较", "部署对齐滑窗模型"],
        [
            ["Epochs", "每折 300，关闭早停", "每折 300，关闭早停"],
            ["Batch size", "16", "64"],
            ["优化器", "AdamW", "AdamW"],
            ["学习率 / Weight decay", "3e-4 / 1e-3", "3e-4 / 1e-3"],
            ["Dropout", "0.5", "0.5"],
            ["模型保存", "验证集 BA 最优", "验证集 BA 最优"],
            ["测试指标", "Accuracy、Precision、Recall、Specificity、F1、BA、混淆矩阵", "同左，并增加整视频事件测试"],
        ],
        [2600, 3380, 3380],
    )
    add_callout(
        doc,
        "训练量说明：",
        "23 路 × 4 折 × 300 轮 = 27,600 epochs，共保存 92 个折模型；"
        "部署对齐模型另训练 4 折 × 300 轮。所有正式结果均来自实际运行产物，而非文献表格。",
        fill=LIGHT,
        accent=DARK_BLUE,
    )

    add_heading(doc, "3. 模型运行结果", 1)
    add_heading(doc, "3.1 GMDCSA24 主体隔离内部测试", 2)
    add_paragraph(
        doc,
        "内部测试中，YOLO-Pose + ByteTrack + ST-GCN++ 的 Balanced Accuracy 最高；"
        "RTMPose + ST-GCN++ 排名第二，但 Specificity 更高，且后续外部测试更稳定。"
    )
    add_table(
        doc,
        ["排名", "路线", "BA", "F1", "Recall", "Specificity"],
        [
            ["1", "YOLO-Pose + ByteTrack + ST-GCN++", "89.41%", "89.57%", "92.41%", "86.42%"],
            ["2", "RTMPose + ST-GCN++", "86.83%", "86.27%", "83.54%", "90.12%"],
            ["3", "RTMPose + ByteTrack + ST-GCN++", "86.29%", "86.59%", "89.87%", "82.72%"],
            ["3", "RTMPose + PoseC3D-style", "86.29%", "86.59%", "89.87%", "82.72%"],
            ["5", "RTMPose + ByteTrack v2 + ST-GCN++", "84.32%", "83.44%", "79.75%", "88.89%"],
            ["6", "RTMPose + CTR-GCN", "83.15%", "83.23%", "84.81%", "81.48%"],
        ],
        [700, 3800, 1215, 1215, 1215, 1215],
        numeric_cols=(0, 2, 3, 4, 5),
    )
    add_picture_with_alt(
        doc, bar_chart, Inches(6.25),
        "GMDCSA24主体隔离测试领先路线Balanced Accuracy横向柱状图",
    )
    cap = add_paragraph(doc, "图 1  GMDCSA24 主体隔离测试中的领先路线（BA）",
                        align=WD_ALIGN_PARAGRAPH.CENTER, before=3, after=8, color=MUTED)
    set_run_font(cap.runs[0], size=9.5, italic=True, color=MUTED)

    add_heading(doc, "3.2 MCFD 跨视角外部测试", 2)
    add_paragraph(
        doc,
        "MCFD 未参与训练。固定阈值 0.5 的跨视角测试显示：内部冠军并不一定是外部最稳健路线。"
        "RTMPose + ST-GCN++ 的整体均衡性较好，而 YOLO-Pose + ST-GCN++ 的召回率和 F1 略高。"
    )
    add_table(
        doc,
        ["路线", "Accuracy", "Recall", "F1", "ROC-AUC"],
        [
            ["RTMPose + ST-GCN++", "62.41%", "60.30%", "60.61%", "64.39%"],
            ["YOLO-Pose + ST-GCN++", "60.00%", "66.33%", "61.40%", "64.84%"],
            ["YOLO-Pose + CTR-GCN", "58.55%", "62.81%", "59.24%", "61.38%"],
        ],
        [4080, 1320, 1320, 1320, 1320],
        numeric_cols=(1, 2, 3, 4),
    )
    add_callout(
        doc,
        "解释：",
        "MCFD 性能明显低于 GMDCSA24，反映跨数据集、跨视角、摄像机位置和跌倒定义差异。"
        "因此不能只依据内部 Accuracy 或单一数据集冠军决定系统模型。",
        fill="FFF8E8",
        accent=GOLD,
    )

    add_heading(doc, "3.3 ByteTrack 消融结论", 2)
    add_bullet(doc, "YOLO-Pose + ByteTrack + ST-GCN++ 获得最高内部 BA，但外部测试约 61.75% BA，未超过稳健基线。")
    add_bullet(doc, "RTMPose 加 ByteTrack 后 Recall 从 83.54% 提高到 89.87%，但 Specificity 从 90.12% 降到 82.72%。")
    add_bullet(doc, "ByteTrack v2 将全零骨架帧由 1518 降为 0，但 BA 仍只有 84.32%。")
    add_bullet(doc, "原因是主轨姿态与回接姿态逐帧切换会形成时序不连续；减少缺失不等于提高动作分类质量。")

    add_heading(doc, "4. 部署对齐滑窗模型", 1)
    add_heading(doc, "4.1 为什么重新训练", 2)
    add_paragraph(
        doc,
        "早期模型在完整视频中均匀抽取 64 帧训练，而实际系统使用连续 64 帧滑动窗口，"
        "两者存在输入分布差异。旧权重在固定 8 段整视频测试中出现 1 次漏检和 1 次误报。"
        "因此，本阶段依据 GMDCSA24 的时间标注重新生成连续窗口，使训练输入与部署输入一致。"
    )
    add_heading(doc, "4.2 数据与正式指标", 2)
    add_bullet(doc, "对 160 段视频的全部 34,172 帧逐帧提取 RTMPose；仅 3 帧无姿态，零姿态率约 0.0088%。")
    add_bullet(doc, "窗口长度 64 帧、步长 16 帧；与跌倒区间重叠不少于 0.5 秒标为正类。")
    add_bullet(doc, "四折最佳轮次分别为 62、95、69、86；虽然完整训练 300 轮，部署加载各折验证最优权重。")
    add_table(
        doc,
        ["Accuracy", "Precision", "Recall", "Specificity", "F1", "Balanced Accuracy"],
        [["84.56%", "72.12%", "81.34%", "85.99%", "76.45%", "83.67%"]],
        [1560, 1560, 1560, 1560, 1560, 1560],
        numeric_cols=(0, 1, 2, 3, 4, 5),
    )
    add_paragraph(
        doc,
        "合并混淆矩阵：TP=388、TN=921、FP=150、FN=89，共 1548 个受试者隔离测试窗口。",
        color=MUTED,
    )
    curve = PROJECT / "results" / "sliding_window_e300_b64" / "rtmpose_stgcnpp" / "learning_curves.png"
    if curve.exists():
        add_picture_with_alt(
            doc, curve, Inches(6.25),
            "RTMPose加ST-GCN++部署对齐模型四折300轮学习曲线",
        )
        cap = add_paragraph(doc, "图 2  部署对齐模型四折 300 轮学习曲线及最佳轮次",
                            align=WD_ALIGN_PARAGRAPH.CENTER, before=3, after=8, color=MUTED)
        set_run_font(cap.runs[0], size=9.5, italic=True, color=MUTED)

    add_heading(doc, "4.3 整视频系统复测", 2)
    add_table(
        doc,
        ["类别", "固定片段", "旧权重结果", "滑窗权重结果"],
        [
            ["Fall", "S1/01、S2/12、S3/05、S4/08", "3/4 检出，1 次漏检", "4/4 检出"],
            ["ADL", "S1/01、S2/01、S3/01、S4/01", "3/4 正确，1 次误报", "4/4 无报警"],
        ],
        [1200, 3540, 2310, 2310],
    )
    add_callout(
        doc,
        "评价边界：",
        "4/4 与 0/4 是端到端流水线冒烟测试，不是无偏准确率。四折集成中的部分模型见过对应 Subject；"
        "正式性能仍以主体隔离测试为准，设备到位后还需扩大事件级现场测试。",
        fill="FCECEC",
        accent=RED,
    )

    add_heading(doc, "5. 最终模型选择与理由", 1)
    add_table(
        doc,
        ["路线", "优势", "不足", "当前定位"],
        [
            ["RTMPose + ST-GCN++", "内部与外部表现均衡；姿态稳定；已完成滑窗对齐和系统演示", "外部泛化仍有限；逐帧姿态计算较重", "系统默认基线"],
            ["YOLO-Pose + ByteTrack + ST-GCN++", "内部 BA 最高、Recall 高", "外部优势不明显；跟踪缺失存在类别偏置", "高分研究分支"],
            ["YOLO-Pose + ST-GCN++", "外部 Recall 和 F1 较好；部署生态成熟", "内部综合指标低于 RTMPose 路线", "少漏报候选/融合候选"],
            ["RTMPose + CTR-GCN", "图卷积机制不同，可作为正交对照", "整体没有超过 ST-GCN++", "保留对照，不作为默认"],
        ],
        [2160, 3240, 2520, 1440],
    )
    add_callout(
        doc,
        "推荐决策：",
        "以 RTMPose + ST-GCN++ 作为当前系统基线；保持统一接口，使后续可替换为 YOLO-Pose、"
        "加入跟踪或进行双路线融合，而不修改上层业务系统。",
        fill="EAF6EF",
        accent=GREEN,
    )

    add_heading(doc, "6. 系统对接方案", 1)
    add_paragraph(
        doc,
        "建议把跌倒检测封装为独立算法服务。业务系统只负责提供摄像头或视频流，并接收标准化报警事件；"
        "姿态模型、时序分类器和阈值策略均留在算法服务内部。"
    )
    add_picture_with_alt(
        doc, architecture, Inches(6.25),
        "摄像头到RTMPose、滑动窗口、ST-GCN++与事件接口的系统对接流程图",
    )
    cap = add_paragraph(doc, "图 3  推荐的算法服务对接流程",
                        align=WD_ALIGN_PARAGRAPH.CENTER, before=3, after=8, color=MUTED)
    set_run_font(cap.runs[0], size=9.5, italic=True, color=MUTED)

    add_heading(doc, "6.1 输入与输出接口", 2)
    add_table(
        doc,
        ["接口类型", "建议内容"],
        [
            ["输入", "RTSP 地址、USB 摄像头、视频文件，或带时间戳的连续图像帧"],
            ["状态接口", "camera_id、运行状态、FPS、姿态有效率、模型版本"],
            ["窗口结果", "track_id、窗口时间、跌倒概率、四折概率、状态"],
            ["确认事件", "event_id、camera_id、track_id、确认时间、概率、证据视频地址"],
            ["推送方式", "第一版 REST + WebSocket；后续可扩展 Webhook、MQTT 或消息队列"],
        ],
        [2040, 7320],
    )

    add_heading(doc, "6.2 接入前必须补充的工程能力", 2)
    for text in [
        "将当前 run_video() 重构为可持续调用的 process_frame()，支持实时逐帧输入。",
        "为每个 camera_id + track_id 保存独立的关键点缓存、64 帧窗口和报警状态机。",
        "增加 USB/RTSP 视频源、断流重连、超时与健康检查。",
        "增加多人跟踪，防止不同人物骨架串轨，并明确遮挡时输出 UNKNOWN。",
        "保存报警前 5 秒和后 10 秒视频证据，提供事件查询与回放。",
        "增加模型版本、日志、配置管理，以及 Docker 或 Windows 服务部署。",
    ]:
        add_number(doc, text)

    add_heading(doc, "7. 需要与教师和系统组共同讨论的问题", 1)
    add_table(
        doc,
        ["讨论项", "建议初始方案", "需要确认"],
        [
            ["使用场景", "室内单摄像头、固定视角、单人优先", "真实房间、安装高度、视野和光照"],
            ["报警目标", "优先控制漏报，同时限制连续误报", "可接受的每小时误报数和报警延迟"],
            ["系统接口", "REST + WebSocket，CONFIRMED 后推送", "上层平台技术栈和事件字段"],
            ["计算位置", "先在带 NVIDIA GPU 的边缘电脑运行", "设备型号、显存、功耗和并发摄像头数"],
            ["隐私要求", "默认只上传事件和短证据片段", "是否允许保存原视频、保存时长和访问权限"],
            ["模型发布", "生产模型版本化，权重与配置绑定", "升级审批、回滚和现场再训练流程"],
        ],
        [1800, 3780, 3780],
    )

    add_heading(doc, "8. 下一阶段计划与验收指标", 1)
    add_table(
        doc,
        ["阶段", "主要任务", "验收输出"],
        [
            ["A. 可复现发布", "发布权重、自动下载脚本、示例视频和一键运行命令", "新电脑可完成演示"],
            ["B. 实时化", "USB/RTSP、process_frame、多路状态管理、报警 API", "连续运行与接口联调"],
            ["C. 现场采集", "不同人员、视角、光照、遮挡、相似 ADL", "独立现场验证集"],
            ["D. 事件级评估", "事件 Recall、每小时误报、检测延迟、姿态失败率、端到端 FPS", "现场测试报告"],
            ["E. 系统交付", "事件证据、日志、权限、部署、升级与回滚", "可部署软件包和接口文档"],
        ],
        [1800, 4680, 2880],
    )
    add_callout(
        doc,
        "建议验收重点：",
        "设备到位后以事件级 Recall、每小时误报数、平均/95 分位报警延迟、姿态失败率和端到端 FPS "
        "作为主要指标，而不是继续只比较离线窗口 Accuracy。",
        fill=LIGHT,
        accent=DARK_BLUE,
    )

    add_heading(doc, "9. 总结", 1)
    add_paragraph(
        doc,
        "本阶段已经完成教师提出的“模型必须实际运行并反馈效果”的要求。项目不仅运行了最初三条候选路线，"
        "还在统一协议下完成 23 条路线的四折 300 轮比较、MCFD 外部测试、ByteTrack 消融、"
        "部署对齐滑窗重训以及整视频系统演示。实验表明，内部最高分路线并不一定具有最佳跨数据集稳定性；"
        "综合考虑精度、特异度、姿态稳定性、外部测试和工程可接入性，RTMPose + ST-GCN++ "
        "是当前最适合作为系统基线的选择。后续应围绕真实设备、多人场景、报警接口和现场事件级指标推进，"
        "再由教师、算法组和系统组共同确定最终上线配置。"
    )

    add_heading(doc, "附录：主要结果与演示文件", 1)
    add_table(
        doc,
        ["内容", "项目内位置"],
        [
            ["23 路模型综合比较", "docs/MODEL_ROUTE_COMPARISON.md"],
            ["全部路线全部指标", "docs/ALL_ROUTE_METRICS.md"],
            ["部署对齐滑窗训练", "docs/SLIDING_WINDOW_TRAINING.md"],
            ["系统设计与接口", "docs/SYSTEM_DESIGN.md"],
            ["内部汇总结果", "results/benchmark_e300_full_summary.csv"],
            ["滑窗汇总结果", "results/sliding_window_e300_b64_summary.csv"],
            ["跌倒演示视频", "outputs/final_sliding_system/fall_subject1_01/annotated.mp4"],
            ["ADL 演示视频", "outputs/final_sliding_system/adl_subject1_01/annotated.mp4"],
        ],
        [2880, 6480],
    )

    doc.core_properties.title = "视觉跌倒检测模型运行测试与系统对接阶段报告"
    doc.core_properties.subject = "模型实测、效果反馈与系统对接"
    doc.core_properties.author = "项目组"
    doc.core_properties.keywords = "fall detection, RTMPose, ST-GCN++, YOLO-Pose, ByteTrack"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
