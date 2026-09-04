from pathlib import Path
import math

import numpy as np
from PIL import Image, ImageDraw, ImageFont

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "assets"
OUT = ROOT / "跌倒检测与预测技术路线演进及指标对比.docx"

NAVY = "183B56"
BLUE = "2878B5"
CYAN = "39A7C7"
TEAL = "2A9D8F"
GREEN = "4D9F65"
ORANGE = "E89A3D"
RED = "C65353"
LIGHT = "F3F7FA"
MID = "DCEAF3"
GRAY = "667480"
WHITE = "FFFFFF"
BLACK = "1D2730"


FONT = "Microsoft YaHei"
FONT_PATH = Path("C:/Windows/Fonts/msyh.ttc")
FONT_BOLD_PATH = Path("C:/Windows/Fonts/msyhbd.ttc")


def pil_font(size, bold=False):
    path = FONT_BOLD_PATH if bold and FONT_BOLD_PATH.exists() else FONT_PATH
    return ImageFont.truetype(str(path), size)


def rgb(hex_color):
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color="CAD5DE", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    total = sum(widths_dxa)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run(run, size=10.5, bold=False, color=BLACK, font=FONT):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = rgb(color)


def style_paragraph(p, before=0, after=5, line=1.1, align=None):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align


def add_text(doc, text, size=10.5, bold=False, color=BLACK, after=5, align=None):
    p = doc.add_paragraph()
    style_paragraph(p, after=after, align=align)
    set_run(p.add_run(text), size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run(r, size=16 if level == 1 else 12.5, bold=True, color=NAVY if level == 1 else BLUE)
    return p


def add_callout(doc, label, text, fill=MID, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [9120])
    set_table_borders(table, color=fill, size="2")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    style_paragraph(p, after=0, line=1.12)
    set_run(p.add_run(label + "  "), size=10.5, bold=True, color=accent)
    set_run(p.add_run(text), size=10.5, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths, font_size=8.5, highlight_rows=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, widths)
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        style_paragraph(p, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_run(p.add_run(str(h)), size=font_size, bold=True, color=WHITE)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        fill = None
        if highlight_rows and ridx in highlight_rows:
            fill = highlight_rows[ridx]
        elif ridx % 2 == 1:
            fill = "F7F9FB"
        for cidx, val in enumerate(row):
            cell = cells[cidx]
            if fill:
                set_cell_shading(cell, fill)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            style_paragraph(p, after=0, line=1.05,
                            align=WD_ALIGN_PARAGRAPH.LEFT if cidx == 0 else WD_ALIGN_PARAGRAPH.CENTER)
            set_run(p.add_run(str(val)), size=font_size, bold=(fill == "E7F4EA"), color=BLACK)
    return table


def add_picture(doc, path, width=6.25, caption=None):
    p = doc.add_paragraph()
    style_paragraph(p, after=3, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run().add_picture(str(path), width=Inches(width))
    if caption:
        cp = doc.add_paragraph()
        style_paragraph(cp, after=7, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_run(cp.add_run(caption), size=8.5, color=GRAY)


def rounded_box(draw, box, outline, fill="white", radius=24, width=3, shadow=True):
    x0, y0, x1, y1 = box
    if shadow:
        draw.rounded_rectangle((x0+8, y0+8, x1+8, y1+8), radius=radius, fill="#DCE3E8")
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def center_text(draw, xy, text, font, fill, anchor="mm"):
    draw.text(xy, text, font=font, fill=fill, anchor=anchor, align="center", spacing=8)


def arrow(draw, start, end, color="#7E94A3", width=6):
    draw.line((start, end), fill=color, width=width)
    x1, y1 = end
    x0, y0 = start
    ang = math.atan2(y1-y0, x1-x0)
    size = 17
    p1 = (x1-size*math.cos(ang-.55), y1-size*math.sin(ang-.55))
    p2 = (x1-size*math.cos(ang+.55), y1-size*math.sin(ang+.55))
    draw.polygon([end, p1, p2], fill=color)


def create_evolution_image():
    img = Image.new("RGB", (2100, 1020), "white")
    d = ImageDraw.Draw(img)
    d.text((80, 60), "技术路线的4次关键升级", font=pil_font(52, True), fill="#183B56")
    d.text((80, 132), "从事后识别，演进为提前预警与多模态确认", font=pil_font(28), fill="#667480")
    boxes = [(90, 250, 960, 565), (1140, 250, 2010, 565), (90, 655, 960, 970), (1140, 655, 2010, 970)]
    items = [
        ("基础检测", "YOLO-Pose / RTMPose\nST-GCN++ · Normal / Fall", BLUE),
        ("稳定性优化", "ByteTrack · 连续滑窗\n质量门控 · 多折集成", CYAN),
        ("提前预警", "未来 3s / 2s / 1s\nLOW · MEDIUM · HIGH", ORANGE),
        ("多模态闭环", "Qwen3-VL 异步复核\n解释 · 摘要 · FastAPI", GREEN),
    ]
    for i, (box, item) in enumerate(zip(boxes, items), 1):
        title, sub, color = item
        rounded_box(d, box, "#"+color)
        d.ellipse((box[0]+40, box[1]+36, box[0]+108, box[1]+104), fill="#"+color)
        center_text(d, (box[0]+74, box[1]+70), str(i), pil_font(30, True), "white")
        d.text((box[0]+135, box[1]+40), title, font=pil_font(40, True), fill="#183B56")
        d.multiline_text((box[0]+48, box[1]+145), sub, font=pil_font(29), fill="#52616C", spacing=18)
    arrow(d, (975, 407), (1125, 407))
    arrow(d, (1575, 580), (1575, 640))
    arrow(d, (1125, 812), (975, 812))
    path = ASSETS / "01_技术路线演进.png"
    img.save(path)
    return path


def create_current_flow_image():
    img = Image.new("RGB", (2300, 760), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 45), "当前完整技术路线", font=pil_font(50, True), fill="#183B56")
    titles = ["视频接入", "骨架提取", "时序分析", "风险分级", "多模态复核", "系统输出"]
    subs = ["萤石 / 本地", "RTMPose", "ST-GCN++", "1/2/3秒 + Fall", "Qwen3-VL", "FastAPI"]
    colors = [NAVY, BLUE, CYAN, ORANGE, GREEN, NAVY]
    xs = [50, 430, 810, 1190, 1570, 1950]
    for i, (x, title, sub, color) in enumerate(zip(xs, titles, subs, colors)):
        box = (x, 230, x+300, 545)
        rounded_box(d, box, "#"+color, radius=22)
        d.ellipse((x+125, 185, x+175, 235), fill="#"+color)
        center_text(d, (x+150, 210), str(i+1), pil_font(24, True), "white")
        center_text(d, (x+150, 350), title, pil_font(34, True), "#183B56")
        center_text(d, (x+150, 445), sub, pil_font(25), "#5D6B75")
        if i < 5:
            arrow(d, (x+310, 385), (xs[i+1]-12, 385))
    center_text(d, (1080, 650), "骨架分支实时运行", pil_font(27, True), "#2878B5")
    center_text(d, (1660, 650), "HIGH / FALL 时触发", pil_font(27, True), "#4D9F65")
    path = ASSETS / "02_当前完整路线.png"
    img.save(path)
    return path


def grouped_chart(path, title, panels):
    img = Image.new("RGB", (2200, 850), "white")
    d = ImageDraw.Draw(img)
    center_text(d, (1100, 55), title, pil_font(45, True), "#183B56")
    colors = ["#2878B5", "#E89A3D", "#4D9F65"]
    legends = ["BA", "Recall", "Specificity"]
    for j, (legend, color) in enumerate(zip(legends, colors)):
        x = 760 + j*260
        d.rectangle((x, 100, x+40, 125), fill=color)
        d.text((x+52, 94), legend, font=pil_font(23), fill="#4A5963")
    for pidx, panel in enumerate(panels):
        x0 = 80 + pidx*1080
        y0, y1 = 185, 710
        names, series, subtitle = panel
        d.text((x0+500, 145), subtitle, font=pil_font(27, True), fill="#183B56", anchor="mm")
        d.line((x0+60, y1, x0+1010, y1), fill="#9AA9B4", width=2)
        d.line((x0+60, y0, x0+60, y1), fill="#9AA9B4", width=2)
        for tick in [50, 60, 70, 80, 90, 100]:
            yy = y1 - (tick-50)/(100-50)*(y1-y0)
            d.line((x0+60, yy, x0+1010, yy), fill="#E6EBEF", width=2)
            d.text((x0+52, yy), str(tick), font=pil_font(18), fill="#71808A", anchor="rm")
        group_w = 900/len(names)
        bar_w = min(50, group_w/4)
        for gi, name in enumerate(names):
            center = x0+90+group_w*(gi+.5)
            for si, values in enumerate(series):
                v = values[gi]
                xx0 = center + (si-1)*bar_w - bar_w*.42
                yy = y1 - (v-50)/50*(y1-y0)
                d.rectangle((xx0, yy, xx0+bar_w*.84, y1), fill=colors[si])
            center_text(d, (center, 770), name, pil_font(18), "#4F5D66")
    img.save(path)
    return path


def create_detection_chart():
    panels = [
        (["YOLO+BT\nST-GCN++", "RTMPose\nST-GCN++", "RTMPose\nCTR-GCN"],
         [[89.41, 86.83, 83.15], [92.41, 83.54, 84.81], [86.42, 90.12, 81.48]],
         "内部测试 · GMDCSA24（160段）"),
        (["YOLO-Pose\nST-GCN++", "RTMPose\nCTR-GCN", "RTMPose\nST-GCN++", "YOLO+BT\nST-GCN++"],
         [[64.01, 63.48, 63.08, 61.75], [61.81, 60.30, 61.81, 63.32], [66.20, 66.67, 64.35, 60.19]],
         "外部测试 · MCFD（415段）")
    ]
    return grouped_chart(ASSETS / "03_检测路线指标对比.png", "跌倒检测路线对比", panels)


def vertical_chart(path, title, groups, legends, series, colors, ymin=50):
    img = Image.new("RGB", (1900, 820), "white")
    d = ImageDraw.Draw(img)
    center_text(d, (950, 55), title, pil_font(44, True), "#183B56")
    legend_start = 430
    for j, (legend, color) in enumerate(zip(legends, colors)):
        x = legend_start + j*260
        d.rectangle((x, 105, x+38, 130), fill=color)
        d.text((x+48, 98), legend, font=pil_font(22), fill="#4A5963")
    x0, x1, y0, y1 = 150, 1790, 180, 680
    for tick in range(ymin, 101, 10):
        yy = y1-(tick-ymin)/(100-ymin)*(y1-y0)
        d.line((x0, yy, x1, yy), fill="#E6EBEF", width=2)
        d.text((x0-15, yy), str(tick), font=pil_font(19), fill="#71808A", anchor="rm")
    group_w = (x1-x0)/len(groups)
    bw = min(75, group_w/(len(series)+1))
    for gi, group in enumerate(groups):
        center = x0+group_w*(gi+.5)
        for si, vals in enumerate(series):
            v = vals[gi]
            xx = center+(si-(len(series)-1)/2)*bw-bw*.4
            yy = y1-(v-ymin)/(100-ymin)*(y1-y0)
            d.rectangle((xx, yy, xx+bw*.8, y1), fill=colors[si])
        center_text(d, (center, 745), group, pil_font(25, True), "#4F5D66")
    img.save(path)
    return path


def create_prefall_chart():
    return vertical_chart(
        ASSETS / "04_提前预测指标.png", "PreVFall 1/2/3秒提前预测",
        ["1秒", "2秒", "3秒"], ["BA", "Recall", "F1", "PR-AUC"],
        [[90.84, 88.17, 85.98], [84.08, 79.82, 80.47], [66.61, 68.87, 64.94], [84.10, 83.09, 82.22]],
        ["#2878B5", "#E89A3D", "#39A7C7", "#4D9F65"], ymin=55)


def create_cascade_chart():
    labels = ["Accuracy", "Precision", "Recall", "Specificity", "F1", "BA"]
    skeleton = [90.74, 90.79, 95.83, 80.56, 93.24, 88.19]
    qwen = [96.30, 100.00, 94.44, 100.00, 97.14, 97.22]
    img = Image.new("RGB", (1900, 850), "white")
    d = ImageDraw.Draw(img)
    center_text(d, (950, 55), "Qwen3-VL触发式复核效果", pil_font(44, True), "#183B56")
    for j, (label, color) in enumerate([("骨架1秒HIGH", "#2878B5"), ("Qwen升级确认", "#4D9F65")]):
        x = 610+j*350
        d.rectangle((x, 105, x+42, 132), fill=color)
        d.text((x+55, 98), label, font=pil_font(22), fill="#4A5963")
    x0, x1 = 400, 1770
    for i, label in enumerate(labels):
        cy = 205+i*95
        d.text((x0-25, cy), label, font=pil_font(23), fill="#4F5D66", anchor="rm")
        for val, color, off in [(skeleton[i], "#2878B5", -18), (qwen[i], "#4D9F65", 18)]:
            start = x0
            end = x0+(val-75)/27*(x1-x0)
            d.rounded_rectangle((start, cy+off-12, end, cy+off+12), radius=7, fill=color)
            d.text((min(end+12, 1840), cy+off), f"{val:.2f}", font=pil_font(18), fill="#4F5D66", anchor="lm")
    img.save(ASSETS / "05_Qwen级联指标.png")
    return ASSETS / "05_Qwen级联指标.png"


def set_doc_styles(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(.72)
    sec.bottom_margin = Inches(.72)
    sec.left_margin = Inches(.85)
    sec.right_margin = Inches(.85)
    sec.header_distance = Inches(.35)
    sec.footer_distance = Inches(.38)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(BLACK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, NAVY, 12, 7),
        ("Heading 2", 12.5, BLUE, 9, 5),
    ]:
        st = doc.styles[name]
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = rgb(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(header.add_run("挑战杯 · 行为检测模块"), size=8.5, color=GRAY)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break()
    p.runs[-1]._element.getparent().remove(p.runs[-1]._element)
    doc.add_page_break()


def build_doc():
    evo = create_evolution_image()
    current = create_current_flow_image()
    det = create_detection_chart()
    pre = create_prefall_chart()
    cas = create_cascade_chart()

    doc = Document()
    set_doc_styles(doc)

    # Page 1
    p = doc.add_paragraph()
    style_paragraph(p, before=8, after=3)
    set_run(p.add_run("跌倒检测与预测"), size=24, bold=True, color=NAVY)
    p2 = doc.add_paragraph()
    style_paragraph(p2, after=12)
    set_run(p2.add_run("技术路线演进及指标对比"), size=17, bold=True, color=BLUE)
    add_callout(doc, "核心结论", "由单一跌倒二分类，升级为“提前预警 + 多模态确认 + 系统接口”的完整闭环。", fill="E8F2F7", accent=NAVY)
    add_picture(doc, evo, width=6.45)
    add_text(doc, "演进主线：事后识别 → 稳定跟踪 → 1/2/3秒预测 → Qwen语义复核", size=10,
             bold=True, color=GRAY, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # Page 2
    add_heading(doc, "1  当前完整路线", 1)
    add_picture(doc, current, width=6.55)
    add_callout(doc, "实时主线", "RTMPose提取骨架，ST-GCN++持续完成跌倒检测与1/2/3秒风险预测。", fill="E8F2F7", accent=BLUE)
    add_callout(doc, "触发分支", "仅在HIGH或FALL时调用Qwen3-VL，输出确认结果、原因解释和事件摘要。", fill="E7F4EA", accent=GREEN)
    add_callout(doc, "工程输出", "统一拉流、共享缓存，通过FastAPI接入护理站或师兄业务系统。", fill="FFF1DE", accent=ORANGE)
    add_text(doc, "设计原则：大模型不阻塞实时预警，也不能撤销骨架告警。", size=11, bold=True,
             color=RED, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # Page 3
    add_heading(doc, "2  跌倒检测路线对比", 1)
    add_text(doc, "内部测试衡量同数据源新受试者；外部测试衡量跨数据集、跨视角泛化。", size=9.5, color=GRAY)
    add_picture(doc, det, width=6.55)
    add_table(doc,
              ["测试口径", "推荐/领先路线", "BA", "Recall", "Specificity", "结论"],
              [
                  ("GMDCSA24内部", "YOLO+ByteTrack+ST-GCN++", "89.41%", "92.41%", "86.42%", "内部冠军"),
                  ("GMDCSA24内部", "RTMPose+ST-GCN++", "86.83%", "83.54%", "90.12%", "默认主线"),
                  ("MCFD外部", "YOLO-Pose+ST-GCN++", "64.01%", "61.81%", "66.20%", "外测BA最高"),
                  ("MCFD外部", "RTMPose+ST-GCN++", "63.08%", "61.81%", "64.35%", "综合稳定"),
              ],
              [1450, 3100, 900, 980, 1120, 1690], font_size=8.0,
              highlight_rows={1: "E8F2F7"})
    add_callout(doc, "路线选择", "ByteTrack内部增益明显，但外部泛化不稳定；系统默认仍采用RTMPose + ST-GCN++。", fill="FFF1DE", accent=ORANGE)

    doc.add_page_break()

    # Page 4
    add_heading(doc, "3  跌倒提前预测", 1)
    add_text(doc, "PreVFall：108段视频、9名受试者、72段跌倒；严格LOSO，输入不包含跌倒发生后的画面。", size=9.5, color=GRAY)
    add_picture(doc, pre, width=6.15)
    add_table(doc,
              ["预测时距", "Balanced Accuracy", "Precision", "Recall", "F1", "PR-AUC", "风险等级"],
              [
                  ("1秒", "90.84%", "62.00%", "84.08%", "66.61%", "84.10%", "HIGH"),
                  ("2秒", "88.17%", "63.81%", "79.82%", "68.87%", "83.09%", "MEDIUM"),
                  ("3秒", "85.98%", "59.14%", "80.47%", "64.94%", "82.22%", "LOW"),
              ],
              [1050, 1580, 1120, 1080, 900, 1050, 1420], font_size=8.4,
              highlight_rows={0: "E7F4EA"})
    add_callout(doc, "结果判断", "1秒预测最可靠；2秒用于中等级提示；3秒仅作为低等级风险，不直接触发最高告警。", fill="E8F2F7", accent=NAVY)

    doc.add_page_break()

    # Page 5
    add_heading(doc, "4  Qwen3-VL多模态级联效果", 1)
    add_text(doc, "协议：PreVFall 108段；骨架1秒HIGH立即预警，缓存2.5秒后由Qwen复核触发前后约4.5秒RGB。", size=9.5, color=GRAY)
    add_picture(doc, cas, width=6.2)
    add_table(doc,
              ["阶段", "TP/TN/FP/FN", "Accuracy", "Precision", "Recall", "Specificity", "F1", "BA"],
              [
                  ("骨架1秒HIGH", "69/29/7/3", "90.74%", "90.79%", "95.83%", "80.56%", "93.24%", "88.19%"),
                  ("Qwen升级确认", "68/36/0/4", "96.30%", "100.00%", "94.44%", "100.00%", "97.14%", "97.22%"),
              ],
              [1550, 1450, 1080, 1100, 920, 1180, 860, 900], font_size=7.8,
              highlight_rows={1: "E7F4EA"})
    add_callout(doc, "提升点", "Qwen拒绝7/7个骨架误触发，系统Accuracy由90.74%提升至96.30%，F1提升至97.14%。", fill="E7F4EA", accent=GREEN)
    add_callout(doc, "边界", "该结果属于触发式离线级联评估；真实家庭长期误报率、遮挡和多人场景仍需现场验证。", fill="FCEAEA", accent=RED)

    # Metadata
    props = doc.core_properties
    props.title = "跌倒检测与预测技术路线演进及指标对比"
    props.subject = "挑战杯行为检测模块答辩材料"
    props.keywords = "跌倒检测, 跌倒预测, RTMPose, ST-GCN++, Qwen3-VL"
    props.author = "项目团队"

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
